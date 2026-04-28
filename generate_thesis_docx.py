from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


BASE_DIR = Path(r"G:\Java-lib")
SOURCE = BASE_DIR / "论文正文.md"
OUTPUT = BASE_DIR / "论文.docx"


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_common(paragraph, first_line_indent=0, before=0, after=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Pt(first_line_indent)


def style_paragraph(paragraph, kind="body"):
    if kind == "body":
        set_paragraph_common(paragraph, first_line_indent=24)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 12, False)
    elif kind == "eng_body":
        set_paragraph_common(paragraph, first_line_indent=24)
        for run in paragraph.runs:
            set_run_font(run, "Times New Roman", "Times New Roman", 12, False)
    elif kind == "h1":
        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=12)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 16, True)
    elif kind == "h2":
        set_paragraph_common(paragraph, before=12, after=6)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 14, True)
    elif kind == "h3":
        set_paragraph_common(paragraph, before=6, after=6)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 12, True)
    elif kind == "title_cn":
        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 16, True)
    elif kind == "title_en":
        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
        for run in paragraph.runs:
            set_run_font(run, "Times New Roman", "Times New Roman", 16, True)
    elif kind == "keyword":
        set_paragraph_common(paragraph)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 12, False)
    elif kind == "reference":
        set_paragraph_common(paragraph)
        fmt = paragraph.paragraph_format
        fmt.hanging_indent = Pt(24)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 10.5, False)
    elif kind == "toc_title":
        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 15, True)
    elif kind == "cover_big":
        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 24, True)
    elif kind == "cover_title":
        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 22, True)
    elif kind == "cover_info":
        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 16, False)
    elif kind == "table":
        set_paragraph_common(paragraph)
        for run in paragraph.runs:
            set_run_font(run, "宋体", "Times New Roman", 10.5, False)


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)
    style_paragraph(paragraph, "keyword")


def insert_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "请在Word中右键目录并选择“更新域”。"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_text)
    run._r.append(fld_char_end)
    style_paragraph(paragraph, "keyword")


def split_blocks(text):
    blocks = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("#"):
            blocks.append(("heading", line))
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue
        para_lines = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("|"):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append(("paragraph", " ".join(para_lines)))
    return blocks


def parse_table(lines):
    rows = []
    for idx, line in enumerate(lines):
        cols = [c.strip() for c in line.strip().strip("|").split("|")]
        if idx == 1 and all(set(c) <= {"-", ":"} for c in cols):
            continue
        rows.append(cols)
    return rows


def add_table(document, rows):
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                style_paragraph(p, "table")
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
    document.add_paragraph()


def add_cover(document):
    p = document.add_paragraph()
    p.add_run("本科毕业设计（论文）")
    style_paragraph(p, "cover_big")

    for _ in range(4):
        document.add_paragraph()

    p = document.add_paragraph()
    p.add_run("题目：基于Vue的考研图书馆座位预约系统的设计与实现")
    style_paragraph(p, "cover_title")

    for _ in range(5):
        document.add_paragraph()

    items = [
        "院（系）：[待补充]",
        "专    业：计算机科学与技术",
        "班    级：[待补充]",
        "学    生：[待补充]",
        "学    号：[待补充]",
        "指导教师：[待补充]",
        "完成时间：2026年4月",
    ]
    for item in items:
        p = document.add_paragraph()
        p.add_run(item)
        style_paragraph(p, "cover_info")


def add_blocks(document, blocks):
    english_mode = False
    first_main = True
    for kind, content in blocks:
        if kind == "heading":
            level = len(content) - len(content.lstrip("#"))
            text = content[level:].strip()

            if text == "Abstract":
                english_mode = True
            elif text.startswith("1 绪论"):
                english_mode = False

            if level == 1 and text not in {"中文摘要", "Abstract", "参考文献", "致谢"} and text.startswith(tuple(str(i) for i in range(1, 8))):
                if not first_main:
                    document.add_page_break()
                first_main = False

            p = document.add_paragraph()
            p.add_run(text)
            if text == "中文摘要":
                style_paragraph(p, "title_cn")
            elif text == "Abstract":
                style_paragraph(p, "title_en")
            elif level == 1:
                style_paragraph(p, "h1")
            elif level == 2:
                style_paragraph(p, "h2")
            else:
                style_paragraph(p, "h3")
        elif kind == "paragraph":
            p = document.add_paragraph()
            p.add_run(content)
            if content.startswith("关键词："):
                style_paragraph(p, "keyword")
            elif content.startswith("Key Words:"):
                style_paragraph(p, "eng_body")
            elif re.match(r"^\[\d+\]", content):
                style_paragraph(p, "reference")
            else:
                style_paragraph(p, "eng_body" if english_mode else "body")
        elif kind == "table":
            add_table(document, parse_table(content))


def main():
    text = SOURCE.read_text(encoding="utf-8")
    blocks = split_blocks(text)

    # Split front matter from body
    body_start = next(i for i, b in enumerate(blocks) if b[0] == "heading" and b[1].startswith("# 1 绪论"))
    front_blocks = blocks[:body_start]
    body_blocks = blocks[body_start:]

    document = Document()
    configure_page(document.sections[0])
    add_cover(document)

    front = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(front)
    add_page_number(front.footer.paragraphs[0])

    title = document.add_paragraph()
    title.add_run("基于Vue的考研图书馆座位预约系统的设计与实现")
    style_paragraph(title, "title_cn")
    add_blocks(document, front_blocks)

    document.add_page_break()
    p = document.add_paragraph()
    p.add_run("目  录")
    style_paragraph(p, "toc_title")
    toc = document.add_paragraph()
    insert_toc(toc)

    body = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body)
    body.header.paragraphs[0].text = "西安工业大学毕业设计（论文）"
    body.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(body.header.paragraphs[0], "table")
    add_page_number(body.footer.paragraphs[0])

    add_blocks(document, body_blocks)

    document.save(OUTPUT)


if __name__ == "__main__":
    main()
